import os
import hashlib
import pickle
from langchain_community.document_loaders import TextLoader, PyPDFLoader
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from .config import logger

# Root and data directories
ROOT_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DATA_DIR = os.path.join(ROOT_DIR, "data")


def get_document_hash(file_path):
    """Calculate MD5 hash of a file to detect changes."""
    with open(file_path, "rb") as f:
        return hashlib.md5(f.read()).hexdigest()


def should_update_embeddings(hash_path):
    """Return True if documents have changed since last embedding."""
    if not os.path.exists(hash_path):
        return True
    try:
        with open(hash_path, "rb") as f:
            stored_hashes = pickle.load(f)
    except Exception:
        return True

    current_files = set(os.listdir(DATA_DIR))
    stored_files = set(stored_hashes.keys())

    if current_files != stored_files:
        return True

    for file in current_files:
        file_path = os.path.join(DATA_DIR, file)
        if stored_hashes.get(file) != get_document_hash(file_path):
            return True

    return False


def _load_docx(file_path):
    """Load a .docx file using python-docx (lightweight, no unstructured needed)."""
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return [Document(page_content=text, metadata={"source": file_path})]
    except Exception as e:
        logger.error(f"Error loading DOCX {file_path}: {e}")
        return []


def load_documents(data_dir=None):
    """Load and chunk all documents from the data directory."""
    if data_dir is None:
        data_dir = DATA_DIR

    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
        logger.error("Data directory created but empty — add documents.")
        return None

    documents = []
    for file in os.listdir(data_dir):
        file_path = os.path.join(data_dir, file)
        try:
            if file.endswith(".pdf"):
                loader = PyPDFLoader(file_path)
                documents.extend(loader.load())
                logger.info(f"Loaded PDF: {file}")
            elif file.endswith(".txt"):
                loader = TextLoader(file_path, encoding="utf-8")
                documents.extend(loader.load())
                logger.info(f"Loaded TXT: {file}")
            elif file.endswith(".docx"):
                docs = _load_docx(file_path)
                documents.extend(docs)
                logger.info(f"Loaded DOCX: {file}")
        except Exception as e:
            logger.error(f"Error loading {file}: {e}")
            continue

    if not documents:
        logger.error("No valid documents found.")
        return None

    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    chunks = splitter.split_documents(documents)
    logger.info(f"Split into {len(chunks)} chunks.")
    return chunks


def save_document_hashes(hash_path, data_dir=None):
    """Persist document hashes for change detection."""
    if data_dir is None:
        data_dir = DATA_DIR

    doc_hashes = {}
    if os.path.exists(hash_path):
        try:
            with open(hash_path, "rb") as f:
                doc_hashes = pickle.load(f)
        except Exception:
            pass

    for file in os.listdir(data_dir):
        doc_hashes[file] = get_document_hash(os.path.join(data_dir, file))

    with open(hash_path, "wb") as f:
        pickle.dump(doc_hashes, f)
